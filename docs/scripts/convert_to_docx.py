import os
import re
import docx
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, hex_color):
    """Sets the background color of a cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_table_borders(table):
    """Sets clean light grey borders to the table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

def add_page_number(run):
    """Adds a dynamic PAGE field to a run in the header."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_runs_from_markdown(paragraph, text):
    """Parses standard inline Markdown elements (**bold**, *italic*, $math$, links) and adds them as runs."""
    # Find markdown blocks
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|\$.*?\$|https?://\S+)')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        
        run = paragraph.add_run()
        # Preserve font name/size by applying to run
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        
        if part.startswith('***') and part.endswith('***'):
            run.text = part[3:-3]
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith('$') and part.endswith('$'):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith('http://') or part.startswith('https://'):
            run.text = part
            run.underline = True
            run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
        else:
            run.text = part

def convert_md_to_docx(md_path, docx_path):
    # Read Markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    doc = Document()
    
    # Configure page settings (A4 size, margins)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    
    # Configure dynamic page number in header for subsequent pages
    section.different_first_page_header_footer = True
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run()
    header_run.font.name = 'Times New Roman'
    header_run.font.size = Pt(10)
    add_page_number(header_run)
    
    # Default Normal Style modifications
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(13)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Heading Styles modifications
    for level, size in [(1, 16), (2, 14), (3, 13)]:
        style_name = f'Heading {level}'
        h_style = doc.styles[style_name]
        h_font = h_style.font
        h_font.name = 'Times New Roman'
        h_font.size = Pt(size)
        h_font.bold = True
        h_font.color.rgb = docx.shared.RGBColor(0, 0, 0)
        h_style.paragraph_format.space_before = Pt(12)
        h_style.paragraph_format.space_after = Pt(6)
        h_style.paragraph_format.line_spacing = 1.5
        
    is_cover_page = True
    in_table = False
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i].replace('\n', '')
        
        # Check for table accumulation
        if line.strip().startswith('|'):
            in_table = True
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # End of table, process the accumulated table lines
            process_table(doc, table_lines)
            table_lines = []
            in_table = False
            # Fall through to process current line
            
        # Check for page break
        if line.strip() == '---' or (line.strip().startswith('---') and len(line.strip()) <= 5):
            doc.add_page_break()
            is_cover_page = False
            i += 1
            continue
            
        # Skip page numbering notes in document (since Word adds them dynamically)
        if line.strip().startswith('*(Trang số thứ tự:'):
            i += 1
            continue
            
        # Skip Mermaid diagrams block
        if line.strip().startswith('```mermaid'):
            # Skip until closing block
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            i += 1
            # Add a clean text indicator about the architecture
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("[Sơ đồ quy trình hệ thống DDQN-ALNS được mô tả chi tiết trong báo cáo gốc]")
            run.italic = True
            run.font.size = Pt(11)
            continue
            
        # Skip standard empty code blocks if any
        if line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("".join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            continue
            
        # Skip HTML breaks
        if line.strip() == '<br>' or line.strip() == '<br/>':
            if is_cover_page:
                doc.add_paragraph() # Add vertical spacing on cover page
            i += 1
            continue
            
        # Check for Headings
        m = re.match(r'^(#{1,6})\s+(.*)$', line.strip())
        if m:
            level = len(m.group(1))
            title_text = m.group(2)
            
            if is_cover_page:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(title_text)
                run.bold = True
                if level == 1:
                    run.font.size = Pt(18)
                elif level == 2:
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(13)
            else:
                p = doc.add_heading(level=min(level, 3))
                p.paragraph_format.keep_with_next = True
                add_runs_from_markdown(p, title_text)
                # Align Heading 1 in centers if appropriate (mostly standard for section titles)
                if level == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
            
        # Check for equations block
        if line.strip().startswith('$$') and line.strip().endswith('$$'):
            eq = line.strip()[2:-2].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(eq)
            run.italic = True
            run.font.size = Pt(13)
            i += 1
            continue
            
        # Process Lists (Bullet points)
        m_list = re.match(r'^(\s*)[-\*]\s+(.*)$', line)
        if m_list:
            indent_spaces = len(m_list.group(1))
            list_text = m_list.group(2)
            level = indent_spaces // 2
            
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            p.paragraph_format.space_after = Pt(3)
            add_runs_from_markdown(p, list_text)
            i += 1
            continue
            
        # Process regular paragraph
        if line.strip():
            p = doc.add_paragraph()
            if is_cover_page:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Set spacing for cover page items
                p.paragraph_format.space_after = Pt(12)
                
            add_runs_from_markdown(p, line.strip())
            
        i += 1
        
    doc.save(docx_path)
    print(f"Document successfully written to: {docx_path}")

def process_table(doc, lines):
    """Parses markdown table lines and adds a stylized table to the document."""
    # Filter out the header-separator row (e.g. | :--- | :--- |)
    table_data = []
    for line in lines:
        if re.search(r'^[|\s:-]+$', line.strip()):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        table_data.append(cells)
        
    if not table_data:
        return
        
    num_rows = len(table_data)
    num_cols = len(table_data[0])
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Process cells
    for r_idx, row in enumerate(table_data):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.width = Inches(6.0 / num_cols) # Distribute evenly across 6 inches
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            
            # Formatting
            if r_idx == 0:
                # Header formatting
                set_cell_background(cell, "F2F2F2") # Light grey header background
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(val)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            else:
                # Body cell
                # Check alignment based on content (numbers centered, text left)
                if val.replace('.','',1).replace('-','',1).replace('+','',1).replace('%','',1).strip().replace(' ','').isalnum():
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_runs_from_markdown(p, val)
                for run in p.runs:
                    run.font.size = Pt(11)

if __name__ == "__main__":
    src_dir = os.path.dirname(os.path.abspath(__file__))
    workspace_dir = os.path.dirname(src_dir)
    md_file = os.path.join(workspace_dir, "Cong_Trinh_Nghien_Cuu_VRPTW.md")
    docx_file = os.path.join(workspace_dir, "Cong_Trinh_Nghien_Cuu_VRPTW.docx")
    
    if os.path.exists(md_file):
        convert_md_to_docx(md_file, docx_file)
    else:
        print(f"Error: {md_file} not found!")
